from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from .puzzle_data import PuzzleData


class Worksheet:
    def __init__(self, puzzle_data: PuzzleData):
        self.width = puzzle_data.width
        self.height = puzzle_data.height
        self.heading = puzzle_data.heading
        self.puzzle = puzzle_data.puzzle
        self.hint: list[str] = puzzle_data.hint
        self.answer = puzzle_data.answer
        self.__configure_settings()

    def __configure_settings(self):
        # write to docx file
        # Write to docx to puzzle.docx
        self.document = Document()
        # changing the page margins
        sections = self.document.sections
        for section in sections:
            # section.top_margin = Cm(1)
            # section.bottom_margin = Cm(0.8)
            section.left_margin = Cm(2.3)
            section.right_margin = Cm(2.3)

    def __add_heading(self, grade: str = "__", class_num: str = "__"):
        head = self.document.add_heading(self.heading, 0)
        head.alignment = WD_ALIGN_PARAGRAPH.CENTER
        belong = f"{grade}학년 {class_num}반 이름: _______"
        para_belong = self.document.add_paragraph(belong)
        para_belong.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    def __add_puzzle(self):
        puzzle_table = self.document.add_table(
            rows=self.height, cols=self.width, style="Table Grid"
        )
        puzzle_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        self.set_height = 7200 / self.height
        for i, row in enumerate(puzzle_table.rows):
            #######################세로 길이 정하기!
            # accessing row xml and setting tr height
            tr = row._tr
            trPr = tr.get_or_add_trPr()
            trHeight = OxmlElement("w:trHeight")
            trHeight.set(qn("w:val"), str(self.set_height))
            trHeight.set(qn("w:hRule"), "atLeast")
            trPr.append(trHeight)

            for j, cell in enumerate(row.cells):
                #####가로 길이 정하기!
                cell.width = Inches(5)
                cell.text = self.puzzle[i][j]
                for paragraph in cell.paragraphs:
                    #####가운데 정렬!!
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    paragraph.style.font.bold = True
                #####상하 방향에서 가운데 정렬
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                tcVAlign = OxmlElement("w:vAlign")
                tcVAlign.set(qn("w:val"), "center")
                tcPr.append(tcVAlign)

    def __add_hint_in_table(self):
        word_num = len(self.hint)
        if word_num <= 15:
            size = 5
        elif word_num <= 21:
            size = (word_num + 2) // 4
        else:
            size = 6
        hint_table = self.document.add_table(
            rows=(word_num + size - 1) // size, cols=size, style="Table Grid"
        )
        hint_table.alignment = WD_TABLE_ALIGNMENT.CENTER

        for i, row in enumerate(hint_table.rows):
            # accessing row xml and setting tr height
            tr = row._tr
            trPr = tr.get_or_add_trPr()
            trHeight = OxmlElement("w:trHeight")
            trHeight.set(qn("w:val"), "60")
            trHeight.set(qn("w:hRule"), "atLeast")
            trPr.append(trHeight)

            for j, cell in enumerate(row.cells):
                index = i * size + j

                # 단어 수 만큼 반복하기
                if index < word_num:
                    for paragraph in cell.paragraphs:
                        run = paragraph.add_run(self.hint[index])
                        font = run.font
                        font.name = "Arial"
                        font.size = Pt(13)

                        #####가운데 정렬!!
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        paragraph.style.font.bold = True
                #####상하 방향에서 가운데 정렬
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                tcVAlign = OxmlElement("w:vAlign")
                tcVAlign.set(qn("w:val"), "center")
                tcPr.append(tcVAlign)

    def write(self):
        self.__add_heading()
        self.__add_puzzle()
        self.document.add_paragraph()
        self.__add_hint_in_table()

    def write_answer(self):
        # 정답 파일 쓰기
        self.document.add_page_break()
        answer_table = self.document.add_table(
            rows=self.height, cols=self.width, style="Table Grid"
        )
        answer_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        for i, row in enumerate(answer_table.rows):
            #######################세로 길이 정하기!
            # accessing row xml and setting tr height
            tr = row._tr
            trPr = tr.get_or_add_trPr()
            trHeight = OxmlElement("w:trHeight")
            trHeight.set(qn("w:val"), str(self.set_height))
            trHeight.set(qn("w:hRule"), "atLeast")
            trPr.append(trHeight)

            for j, cell in enumerate(row.cells):
                #####가로 길이 정하기!
                cell.width = Inches(8)
                cell.text = str(self.answer[i][j])
                for paragraph in cell.paragraphs:
                    #####가운데 정렬!!
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    paragraph.style.font.bold = True
                    if cell.text == "0":
                        for run in paragraph.runs:
                            run.font.color.rgb = RGBColor(255, 255, 255)
                    else:
                        for run in paragraph.runs:
                            run.font.color.rgb = RGBColor(255, 0, 0)
                #####상하 방향에서 가운데 정렬
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                tcVAlign = OxmlElement("w:vAlign")
                tcVAlign.set(qn("w:val"), "center")
                tcPr.append(tcVAlign)

    def save(self, filename_or_stream):
        # TODO get answer of puzzle too
        return self.document.save(filename_or_stream)
